"""Genera el documento Word de la Entrega 3 a partir de la plantilla de la catedra."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = BASE_DIR / "Entrega N° 3_ Evolucion del Sistema Inteligente.docx"

MEMBERS = [
    ("Bodetto, Tomás Agustín", "tbodetto@frba.utn.edu.ar", "14%"),
    ("Jugo, Germán Ignacio", "gjugo@frba.utn.edu.ar", "20%"),
    ("Ledesma, Nicolás Ezequiel", "nicoledesma@frba.utn.edu.ar", "14%"),
    ("Marchan Figuera, Brisabed Alexandra", "bmarchanfiguera@frba.utn.edu.ar", "20%"),
    ("Manfredi, Valentín Nehuen", "vamanfredi@frba.utn.edu.ar", "14%"),
    ("Stazzone, Franco Alejandro", "fstazzone@frba.utn.edu.ar", "18%"),
]
PRESENTATION_DATE = "04/07/2026"

doc = Document()

# ---------- Portada ----------
title_table = doc.add_table(rows=1, cols=2)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
title_table.cell(0, 1).text = "Inteligencia Artificial Avanzada"
title_table.cell(0, 1).paragraphs[0].runs[0].bold = True

doc.add_paragraph()
h = doc.add_heading("Entrega N° 3: Evolución del Prototipo del Sistema Inteligente", level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Grupo B")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
members_table = doc.add_table(rows=1 + len(MEMBERS), cols=3)
members_table.style = "Table Grid"
hdr = members_table.rows[0].cells
hdr[0].text = "Apellido y Nombres"
hdr[1].text = "E-Mail"
hdr[2].text = "Porcentaje de Aporte"
for row, (name, email, pct) in zip(members_table.rows[1:], MEMBERS):
    row.cells[0].text = name
    row.cells[1].text = email
    row.cells[2].text = pct

doc.add_paragraph()
date_table = doc.add_table(rows=1, cols=2)
date_table.style = "Table Grid"
date_table.rows[0].cells[0].text = "Fecha de Presentación"
date_table.rows[0].cells[1].text = PRESENTATION_DATE

doc.add_page_break()

# ---------- Resumen ----------
doc.add_heading("Resumen", level=1)
doc.add_paragraph(
    "Este documento presenta la tercera entrega del trabajo práctico, correspondiente a la evolución "
    "del prototipo de Sistema Inteligente para la predicción del éxito comercial de canciones "
    "(clasificación hit / no hit) a partir de datos de Spotify y del ranking Billboard Hot 100. "
    "A partir de la corrección señalada sobre la Entrega 2 —la existencia de información duplicada en "
    "el dataset de entrenamiento— se identificó que el problema real no era la repetición semanal de una "
    "misma canción en sí (eso es parte de cómo se comporta un hit en la realidad), sino que el corte "
    "cronológico de entrenamiento/validación/prueba permitía que una misma canción quedara repartida "
    "entre esos subconjuntos. Se corrigió el esquema de selección de datos mediante un split por grupo de "
    "canción, y se compararon tres arquitecturas de red híbrida (LSTM, GRU y Conv1D) para la rama temporal "
    "del modelo bajo esta metodología corregida. Se documentan las métricas obtenidas, su comparación con "
    "los resultados de la Entrega 2 y las lecciones aprendidas durante el proceso, incluyendo una primera "
    "iteración de corrección que fue descartada por eliminar información relevante del problema."
)

# ---------- 1. Introducción ----------
doc.add_heading("1. Introducción", level=1)
doc.add_paragraph(
    "El problema abordado por el Sistema Inteligente continúa siendo la identificación temprana de "
    "canciones con potencial de éxito comercial, dado que las decisiones de promoción e inversión en la "
    "industria musical suelen basarse en criterios subjetivos o en análisis manuales de datos. El objetivo "
    "del prototipo es clasificar canciones como \"hit\" o \"no hit\" combinando las características propias "
    "de cada canción con el contexto de las tendencias musicales predominantes en las semanas previas."
)
doc.add_paragraph(
    "En la evaluación de la Entrega 2 se destacó que la red híbrida LSTM + MLP era una arquitectura "
    "adecuada para el problema y que el análisis de resultados estaba bien elaborado, pero se señalaron dos "
    "observaciones concretas a resolver en esta entrega: (1) ajustar la preparación de los datos para "
    "evitar información duplicada, y (2) evaluar el uso de otros tipos de capas para la rama temporal, "
    "como GRU y/o ConvNet (Conv1D). Esta entrega aborda ambos puntos: se corrigió el esquema de selección "
    "de datos para eliminar la fuga de información detectada, y se entrenaron y compararon tres variantes "
    "de arquitectura (LSTM, GRU y Conv1D) bajo las mismas condiciones, evaluando su impacto en las métricas "
    "obtenidas y comparándolas con los resultados del prototipo anterior."
)

# ---------- 2. Materiales Disponibles ----------
doc.add_heading("2. Materiales Disponibles", level=1)
doc.add_paragraph(
    "Se mantienen las mismas dos fuentes de datos utilizadas en la Entrega 2, obtenidas desde Kaggle: "
    "el Spotify Tracks Dataset (características musicales cuantificables por canción) y el dataset "
    "histórico Billboard Hot 100 (posición semanal en el ranking)."
)
doc.add_paragraph(
    "Motivo de acotar el análisis al año 2021: el cruce entre ambos datasets se realiza por nombre de "
    "canción y artista normalizados (clean_name / clean_artists), y el año 2021 fue el período en el que "
    "el Spotify Tracks Dataset y el dataset de Billboard Hot 100 presentaban cobertura simultánea "
    "confiable, es decir, la mayor proporción de canciones del ranking de esas semanas se encontraban "
    "también representadas con sus características de audio en el dataset de Spotify. Acotar el análisis "
    "a ese período permitió maximizar la cantidad de coincidencias válidas en el join y evitar introducir "
    "canciones sin información de audio completa."
)

doc.add_heading("2.1 Corrección de la preparación de los datos: información duplicada", level=2)
doc.add_paragraph(
    "En la Entrega 2, el dataset de entrenamiento se construyó a nivel de \"semana x canción\": una misma "
    "canción que permanecía varias semanas en el ranking Billboard generaba una fila de entrenamiento "
    "distinta por cada semana en la que apareció. De 1000 filas etiquetadas como hit, solo 98 correspondían "
    "a canciones distintas (en promedio, cada canción exitosa permanecía cerca de 10 semanas en el "
    "ranking), mientras que las 1450 canciones no-hit eran todas distintas entre sí (sin repetición). El "
    "problema de fondo no es que existan varias filas por canción —cada semana aporta un contexto de "
    "tendencias musicales distinto, por lo que no son copias idénticas—, sino que el corte de "
    "entrenamiento/validación/prueba de la Entrega 2 era puramente cronológico a nivel de fila: una misma "
    "canción podía tener semanas en entrenamiento y semanas en prueba, filtrando su \"fingerprint\" de audio "
    "(que sí es constante) entre ambos conjuntos."
)
doc.add_paragraph(
    "Primera iteración de corrección (descartada): inicialmente se intentó resolver el problema "
    "eliminando toda repetición, dejando una única fila por canción (su primera semana con contexto "
    "temporal completo). Al revisar el resultado se detectó que este enfoque introducía dos problemas "
    "nuevos: (a) descartaba una señal real del fenómeno estudiado —que un hit permanece varias semanas en "
    "el ranking es información legítima, no ruido—, y (b) redujo drásticamente el tamaño de la clase "
    "positiva, de 769 a apenas 79 canciones hit, dejando un conjunto de prueba con solo 20 casos positivos "
    "y métricas con muy poca robustez estadística. Se concluyó que ese enfoque \"tiraba la esencia del "
    "problema\" para resolver un síntoma, en lugar de atacar la causa real."
)
doc.add_paragraph(
    "Corrección final aplicada: en lugar de eliminar filas, se resolvió el problema en el esquema de "
    "selección de datos. Se mantiene el dataset completo a nivel de \"cancion x semana\" (2019 filas, 769 "
    "hits, igual que en la Entrega 2), y el split de entrenamiento/validación/prueba se realiza por "
    "grupo de canción (clean_name + clean_artists): todas las semanas de una misma canción quedan "
    "garantizadas dentro de un único subconjunto. De esta forma se elimina la fuga de información sin "
    "sacrificar el tamaño de muestra ni la señal temporal real del fenómeno."
)

doc.add_heading("2.2 Selección de datos para entrenamiento, validación y prueba", level=2)
doc.add_paragraph(
    "Dado que las canciones hit permanecen en el ranking una cantidad de semanas muy variable (entre 1 y "
    "25 semanas), un split por grupo aleatorio simple podría desviarse bastante de las proporciones "
    "buscadas si, por azar, quedan agrupadas varias canciones de muchas semanas en un mismo subconjunto. "
    "Para evitar esto se implementó un algoritmo de asignación voraz (greedy bin-packing): las canciones de "
    "cada clase (hit / no-hit) se ordenan aleatoriamente y luego por cantidad de semanas en forma "
    "descendente, y cada canción se asigna íntegramente al subconjunto (entrenamiento, validación o "
    "prueba) que en ese momento tenga mayor déficit respecto de su proporción objetivo de filas. Se usó "
    "una semilla fija (42) y como objetivo 60% entrenamiento / 15% validación / 25% prueba, cumpliendo el "
    "mínimo de 25% para el conjunto de prueba exigido por la cátedra."
)
doc.add_paragraph(
    "El resultado fue: 1212 filas para entrenamiento (462 hits, 786 canciones distintas), 303 filas para "
    "validación (115 hits, 207 canciones) y 504 filas para prueba (192 hits, 336 canciones), lo que "
    "equivale a proporciones de 60,0% / 15,0% / 25,0% del total de filas, prácticamente idénticas al "
    "objetivo. Se verificó de forma explícita que ninguna canción aparece en más de un subconjunto."
)

# ---------- 3. Solución Propuesta ----------
doc.add_heading("3. Solución Propuesta", level=1)
doc.add_paragraph(
    "Se mantiene el enfoque general de la Entrega 2: una red híbrida que combina una rama temporal, "
    "encargada de procesar el contexto de las 4 semanas previas del ranking, con una rama densa (MLP) "
    "que procesa las características propias de la canción candidata. Las representaciones de ambas ramas "
    "se concatenan y se procesan mediante capas densas finales hasta una neurona de salida con activación "
    "sigmoide, que estima la probabilidad de que la canción sea un hit."
)
doc.add_paragraph(
    "Siguiendo la sugerencia del docente, en esta entrega se generalizó la rama temporal para poder "
    "intercambiar el tipo de capa utilizada, y se entrenaron y compararon tres variantes bajo las mismas "
    "condiciones de datos, preprocesamiento e hiperparámetros de entrenamiento:"
)
b = doc.add_paragraph(style="List Bullet")
b.add_run("LSTM (línea base de la Entrega 2): ").bold = True
b.add_run("una capa LSTM de 32 unidades procesa la secuencia de 4 semanas.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("GRU: ").bold = True
b.add_run("se reemplaza la capa LSTM por una capa GRU de 32 unidades, con menos parámetros por no tener "
          "puerta de salida (\"output gate\") independiente, lo que puede reducir el sobreajuste en un "
          "dataset pequeño.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("Conv1D (ConvNet): ").bold = True
b.add_run("se reemplaza la rama recurrente por dos capas Conv1D de 32 filtros y kernel de tamaño 2, "
          "seguidas de un GlobalMaxPooling1D, para capturar patrones locales entre semanas consecutivas "
          "sin mantener un estado recurrente.")
doc.add_paragraph(
    "En los tres casos, la rama de la canción candidata (40 variables luego de codificar key/mode/"
    "time_signature) se procesa con una capa densa de 32 neuronas, se concatena con la salida de la rama "
    "temporal, y se agrega una capa densa de 32 neuronas con Dropout antes de la salida sigmoide."
)

arch_img = BASE_DIR / "figura_arquitectura_entrega3.png"
if arch_img.exists():
    doc.add_picture(str(arch_img), width=Cm(15))
    cap = doc.add_paragraph("Figura 1. Arquitectura conceptual de la Entrega 3: la rama temporal es intercambiable entre LSTM, GRU y Conv1D.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

doc.add_heading("3.1 Herramientas, frameworks y librerías", level=2)
doc.add_paragraph(
    "Se mantiene Python como lenguaje principal y TensorFlow/Keras para la construcción y entrenamiento "
    "de las redes neuronales, junto con Pandas para la manipulación de datos y Scikit-learn para el "
    "preprocesamiento (escalado, codificación one-hot) y el cálculo de métricas. A diferencia de la "
    "Entrega 2, cuyo entrenamiento se documentó para ejecutarse en Google Colab, en esta entrega el "
    "entrenamiento de las tres arquitecturas se ejecutó en un entorno local (Python 3.11 + TensorFlow-CPU "
    "2.21), lo que permitió iterar más rápido sobre las distintas variantes de dataset y de arquitectura."
)

# ---------- 4. Desarrollo ----------
doc.add_heading("4. Desarrollo", level=1)
doc.add_paragraph(
    "El desarrollo de esta entrega se organizó en los siguientes scripts, incluidos en la carpeta "
    "scripts/ de esta entrega:"
)
b = doc.add_paragraph(style="List Bullet")
b.add_run("build_grouped_dataset.py: ").bold = True
b.add_run("reconstruye el dataset completo a nivel de cancion x semana (igual granularidad que la "
          "Entrega 2, con la corrección de parsing/normalización de valores) a partir de "
          "tracks_charts_hit_cleaned.csv y tracks_charts_not_hit_cleaned.csv, generando "
          "tracks_charts_grouped_dataset.csv.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("train_all_architectures_grouped.py: ").bold = True
b.add_run("implementa el split por grupo de canción (greedy bin-packing), entrena, evalúa y guarda los "
          "artefactos (modelo, preprocesadores, umbral óptimo, métricas, predicciones y figuras) de las "
          "tres arquitecturas (LSTM, GRU, Conv1D), además de un archivo split_asignacion_canciones.csv "
          "con el subconjunto asignado a cada canción, para trazabilidad.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("build_dedup_dataset.py / train_all_architectures.py: ").bold = True
b.add_run("corresponden a la primera iteración de corrección (deduplicación total), conservados en la "
          "carpeta de esta entrega a modo de registro del proceso, pero reemplazados por el enfoque de "
          "split por grupo descripto en la sección 2.1.")
doc.add_paragraph(
    "Para cada arquitectura se aplicaron pesos de clase balanceados (class_weight) para compensar el "
    "desbalance entre hits y no-hits, Early Stopping monitoreando el AUC de validación, y selección del "
    "umbral de decisión por maximización del F1-score sobre el conjunto de validación, de la misma manera "
    "que en la Entrega 2."
)
doc.add_paragraph(
    "El principal problema encontrado durante el desarrollo fue el ya descripto en la sección 2.1: la "
    "primera corrección aplicada (deduplicar a una fila por canción) resultaba en un conjunto de prueba "
    "con solo 20 canciones hit, e incluso con un split aleatorio estratificado, un corte cronológico sobre "
    "ese dataset dejaba apenas 1 hit en validación y 2 en test. Esto llevó a reconsiderar el enfoque y "
    "migrar al split por grupo de canción sobre el dataset completo, que se describe en la sección 2.2."
)

# ---------- 5. Resultados ----------
doc.add_heading("5. Resultados", level=1)
doc.add_paragraph(
    "Las tres arquitecturas fueron entrenadas y evaluadas sobre los mismos conjuntos de entrenamiento "
    "(1212 filas, 462 hits, 786 canciones), validación (303 filas, 115 hits, 207 canciones) y prueba "
    "(504 filas, 192 hits, 336 canciones). La siguiente tabla resume las métricas obtenidas sobre el "
    "conjunto de prueba, con el umbral de decisión ajustado por maximización de F1 en el conjunto de "
    "validación:"
)

metrics_rows = [
    ("LSTM", "0,39", "98,02%", "95,50%", "99,48%", "97,45%", "99,80%", "99,68%"),
    ("GRU", "0,84", "96,23%", "98,87%", "91,15%", "94,85%", "99,86%", "99,76%"),
    ("Conv1D", "0,13", "95,04%", "88,48%", "100,00%", "93,89%", "99,70%", "99,45%"),
]
table = doc.add_table(rows=1 + len(metrics_rows), cols=8)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, label in enumerate(["Arquitectura", "Umbral", "Accuracy", "Precision", "Recall", "F1-score", "ROC AUC", "PR AUC"]):
    hdr[i].text = label
    hdr[i].paragraphs[0].runs[0].bold = True
for row, values in zip(table.rows[1:], metrics_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

doc.add_paragraph()
doc.add_paragraph(
    "Matrices de confusión sobre el conjunto de prueba (504 filas, 192 hits reales):"
)
b = doc.add_paragraph(style="List Bullet")
b.add_run("LSTM: ").bold = True
b.add_run("303 verdaderos negativos, 9 falsos positivos, 1 falso negativo, 191 verdaderos positivos.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("GRU: ").bold = True
b.add_run("310 verdaderos negativos, 2 falsos positivos, 17 falsos negativos, 175 verdaderos positivos.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("Conv1D: ").bold = True
b.add_run("287 verdaderos negativos, 25 falsos positivos, 0 falsos negativos, 192 verdaderos positivos.")

for arch, caption in [("lstm", "Figura 2. Matriz de confusión - LSTM (test, split por grupo)."),
                       ("gru", "Figura 3. Matriz de confusión - GRU (test, split por grupo)."),
                       ("conv1d", "Figura 4. Matriz de confusión - Conv1D (test, split por grupo).")]:
    img_path = BASE_DIR / f"resultados_{arch}_grouped" / f"figura_matriz_confusion_{arch}.png"
    if img_path.exists():
        doc.add_picture(str(img_path), width=Cm(9))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True

# ---------- 6. Análisis de los Resultados ----------
doc.add_heading("6. Análisis de los Resultados", level=1)
doc.add_paragraph(
    "La arquitectura LSTM obtuvo el mejor desempeño global sobre el conjunto de prueba (F1 97,45%, "
    "recall 99,48%, ROC AUC 99,80%), seguida por GRU (F1 94,85%, mayor precisión pero menor recall) y "
    "Conv1D (F1 93,89%, recall perfecto pero más falsos positivos). Las tres arquitecturas muestran un "
    "ROC AUC y PR AUC muy altos y cercanos entre sí (por encima de 99,4% en los tres casos), lo que indica "
    "que las tres son capaces de ordenar correctamente las canciones por probabilidad de éxito; las "
    "diferencias más marcadas aparecen en el punto de corte elegido por F1, donde LSTM logra el mejor "
    "equilibrio entre precisión y recall. A diferencia de la comparación realizada en la primera iteración "
    "(con el dataset deduplicado y muy pocos positivos), en esta versión final el conjunto de prueba tiene "
    "192 canciones hit, por lo que las diferencias entre arquitecturas son estadísticamente más confiables. "
    "Se selecciona LSTM como arquitectura recomendada para esta entrega, confirmando la elección original "
    "de la Entrega 2 incluso bajo una evaluación sin fuga de información."
)

doc.add_heading("6.1 Comparación de los datasets (Entrega 2 vs. Entrega 3)", level=2)
doc.add_paragraph(
    "Antes de comparar métricas, es importante comparar los datasets y el criterio de split sobre los que "
    "se entrenó y evaluó cada versión del sistema, ya que la corrección de esta entrega está principalmente "
    "en la metodología de evaluación y no en el volumen de datos:"
)
dataset_rows = [
    ("Unidad de análisis", "Una fila por canción y semana en el ranking", "Una fila por canción y semana en el ranking (sin cambios)"),
    ("Filas totales", "2019", "2019"),
    ("Canciones hit únicas", "98 (representadas en 1000 filas)", "79 con contexto completo (representadas en 769 filas)"),
    ("Proporción de hits (filas)", "38,1%", "38,1% (sin cambios; no se elimina información)"),
    ("Método de split", "Cronológico por fecha, a nivel de fila", "Por grupo de canción (greedy bin-packing), estratificado por clase"),
    ("Tamaño del conjunto de prueba", "252 filas (12,5%, por debajo del mínimo pedido)", "504 filas (25,0%, cumple el mínimo pedido)"),
    ("Riesgo de fuga de información", "Sí: una misma canción podía tener semanas en train y en test", "No: cada canción (con todas sus semanas) queda en un único subconjunto"),
]
dtable = doc.add_table(rows=1 + len(dataset_rows), cols=3)
dtable.style = "Table Grid"
dhdr = dtable.rows[0].cells
for i, label in enumerate(["Aspecto", "Entrega 2", "Entrega 3"]):
    dhdr[i].text = label
    dhdr[i].paragraphs[0].runs[0].bold = True
for row, values in zip(dtable.rows[1:], dataset_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

doc.add_paragraph()
doc.add_paragraph("Ventajas del esquema de datos/split construido en la Entrega 3:").runs[0].bold = True
adv_items = [
    "Elimina la fuga de información entre entrenamiento y prueba sin descartar filas: se resuelve en el "
    "split, no en el dataset.",
    "Conserva el tamaño de muestra completo (2019 filas, 769 hits), evitando el problema de la primera "
    "iteración de corrección, que había reducido la clase positiva a 79 canciones.",
    "Preserva la señal real del fenómeno: que un hit permanece varias semanas en el ranking es información "
    "legítima sobre su éxito, no un artefacto a eliminar.",
    "Cumple el requisito de la cátedra de un mínimo de 25% de los ejemplos para el conjunto de prueba "
    "(la Entrega 2 utilizaba 12,5%), con proporciones finales de 60,0% / 15,0% / 25,0% muy cercanas al "
    "objetivo gracias al algoritmo de asignación voraz.",
    "Permite una comparación de arquitecturas más confiable: con 192 hits en el conjunto de prueba (frente "
    "a 20 en la primera iteración), las diferencias de F1/recall entre LSTM, GRU y Conv1D son "
    "estadísticamente más robustas.",
]
for item in adv_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6.2 Comparación de métricas", level=2)
doc.add_paragraph(
    "La siguiente tabla compara las métricas de prueba de la Entrega 2 (arquitectura LSTM, split "
    "cronológico con riesgo de fuga, 252 filas de prueba) contra las tres arquitecturas de la Entrega 3 "
    "(split por grupo de canción, 504 filas de prueba):"
)
comp_rows = [
    ("Entrega 2 (LSTM, split cronológico con fuga)", "98,81%", "94,55%", "100,00%", "97,20%", "99,69%"),
    ("Entrega 3 - LSTM (split por grupo)", "98,02%", "95,50%", "99,48%", "97,45%", "99,80%"),
    ("Entrega 3 - GRU (split por grupo)", "96,23%", "98,87%", "91,15%", "94,85%", "99,86%"),
    ("Entrega 3 - Conv1D (split por grupo)", "95,04%", "88,48%", "100,00%", "93,89%", "99,70%"),
]
table2 = doc.add_table(rows=1 + len(comp_rows), cols=6)
table2.style = "Table Grid"
hdr2 = table2.rows[0].cells
for i, label in enumerate(["Versión", "Accuracy", "Precision", "Recall", "F1-score", "ROC AUC"]):
    hdr2[i].text = label
    hdr2[i].paragraphs[0].runs[0].bold = True
for row, values in zip(table2.rows[1:], comp_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value

doc.add_paragraph()
img_path = BASE_DIR / "figura_comparacion_entrega2_vs_entrega3.png"
if img_path.exists():
    doc.add_picture(str(img_path), width=Cm(15))
    cap = doc.add_paragraph("Figura 5. Comparación de métricas de test entre la Entrega 2 y las tres arquitecturas de la Entrega 3 (split por grupo).")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True

doc.add_paragraph(
    "El resultado más relevante de esta comparación es que, al eliminar por completo el riesgo de fuga de "
    "información (split por grupo de canción en lugar de split cronológico a nivel de fila), las métricas "
    "de LSTM en la Entrega 3 se mantuvieron prácticamente iguales a las de la Entrega 2, e incluso "
    "mejoraron levemente en F1 (97,45% vs. 97,20%) y en ROC AUC (99,80% vs. 99,69%). Esto es una señal "
    "positiva importante: indica que el buen desempeño observado en la Entrega 2 no era, en su mayor "
    "parte, un artefacto de la fuga de información, sino que el modelo efectivamente aprendió patrones "
    "generalizables entre las características de una canción, el contexto de tendencias musicales y la "
    "probabilidad de éxito. La corrección metodológica de esta entrega no buscaba necesariamente bajar las "
    "métricas, sino garantizar que fueran confiables; el hecho de que se mantengan altas bajo una "
    "evaluación más estricta fortalece la validez del Sistema Inteligente desarrollado."
)

doc.add_heading("6.3 Por qué un split aleatorio simple no alcanza: comparación con una variante descartada", level=2)
doc.add_paragraph(
    "Durante el desarrollo se evaluó también una variante intermedia que usa el dataset completo "
    "(cancion x semana, sin deduplicar, igual que la solución final) pero con una selección de "
    "entrenamiento/validación/prueba mediante train_test_split estratificado por clase a nivel de FILA, "
    "sin agrupar por canción. Esta variante se descartó porque, al asignar cada fila de forma independiente, "
    "vuelve a permitir que una misma canción tenga semanas en más de un subconjunto: se verificó que 56 "
    "canciones quedan repartidas entre entrenamiento y prueba, 43 entre entrenamiento y validación, y 41 "
    "entre validación y prueba (por ejemplo, \"Calling My Phone\" de Lil Tjay y 6LACK aporta 10 filas que "
    "terminan repartidas entre ambos conjuntos). Es decir, reproduce el mismo problema metodológico de la "
    "Entrega 2, solo que sobre un dataset distinto."
)
comp3_rows = [
    ("Entrega 2 (split cronológico, con fuga)", "98,81%", "94,55%", "100,00%", "97,20%", "99,69%"),
    ("Variante descartada (split aleatorio por fila, con fuga)", "99,41%", "98,46%", "100,00%", "99,22%", "99,63%"),
    ("Entrega 3 final (split por grupo de canción, sin fuga)", "98,02%", "95,50%", "99,48%", "97,45%", "99,80%"),
]
table3 = doc.add_table(rows=1 + len(comp3_rows), cols=6)
table3.style = "Table Grid"
hdr3 = table3.rows[0].cells
for i, label in enumerate(["Versión (arquitectura LSTM)", "Accuracy", "Precision", "Recall", "F1-score", "ROC AUC"]):
    hdr3[i].text = label
    hdr3[i].paragraphs[0].runs[0].bold = True
for row, values in zip(table3.rows[1:], comp3_rows):
    for cell, value in zip(row.cells, values):
        cell.text = value
doc.add_paragraph()
doc.add_paragraph(
    "La variante descartada obtiene las métricas más altas de las tres (F1 99,22%), pero por el motivo "
    "equivocado: al filtrar canciones entre entrenamiento y prueba, el modelo puede apoyarse parcialmente "
    "en haber \"visto\" esas canciones durante el entrenamiento. Esto confirma, con un segundo ejemplo "
    "independiente del de la Entrega 2, que un split aleatorio o cronológico a nivel de fila no es "
    "suficiente cuando existen múltiples filas por entidad (en este caso, por canción): es necesario un "
    "split por grupo para que las métricas reportadas sean confiables. Por este motivo se descarta esta "
    "variante y se mantiene el split por grupo de canción (sección 2.2) como solución final de esta "
    "entrega."
)

# ---------- 7. Conclusiones ----------
doc.add_heading("7. Conclusiones", level=1)
doc.add_paragraph(
    "La corrección de la Entrega 2 requirió dos iteraciones: la primera, más agresiva (deduplicar a una "
    "fila por canción), resultó contraproducente porque eliminaba información real del fenómeno y reducía "
    "demasiado la muestra; la segunda, un split por grupo de canción sobre el dataset completo, resolvió "
    "la fuga de información sin sacrificar tamaño de muestra ni señal temporal. Bajo esta metodología "
    "corregida, LSTM se confirma como la arquitectura más adecuada para el problema, superando a GRU y "
    "Conv1D, y sus métricas resultaron prácticamente equivalentes (e incluso levemente superiores en F1 y "
    "ROC AUC) a las reportadas en la Entrega 2, lo que valida que el desempeño del prototipo original era "
    "en gran medida genuino y no un artefacto de la fuga de información. Para futuras iteraciones se "
    "recomienda ampliar el rango temporal de datos (incorporar más años con cobertura simultánea entre "
    "ambos datasets) para robustecer aún más la evaluación, y explorar técnicas adicionales de manejo de "
    "desbalance de clases, como sobremuestreo (SMOTE) o funciones de pérdida focales (focal loss), como "
    "alternativa o complemento a los pesos de clase ya utilizados."
)

# ---------- 8. Lecciones Aprendidas ----------
doc.add_heading("8. Lecciones Aprendidas", level=1)
b = doc.add_paragraph(style="List Bullet")
b.add_run("No toda repetición de datos es \"información duplicada\" en sentido negativo. ").bold = True
b.add_run("Que una misma canción aparezca en varias semanas del ranking es parte real del fenómeno que se "
          "modela (la persistencia de un hit); el problema no era la repetición en sí, sino permitir que "
          "esa repetición cruzara la frontera entre entrenamiento y prueba.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("La primera solución no siempre es la correcta, y conviene revisarla críticamente. ").bold = True
b.add_run("Deduplicar por completo pareció al principio la forma más directa de \"eliminar duplicados\", "
          "pero al analizar el resultado (79 hits en vez de 769, solo 20 en test) se hizo evidente que "
          "resolvía el síntoma tirando información valiosa del problema real. Revisar los resultados "
          "intermedios con sentido crítico evitó quedarse con una solución peor que el problema original.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("El lugar correcto para resolver una fuga de información suele ser el split, no el dataset. ").bold = True
b.add_run("Un split por grupo (por entidad, en este caso la canción) permite conservar todos los datos "
          "disponibles y al mismo tiempo garantizar que ninguna entidad se filtre entre subconjuntos.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("Con grupos de tamaño desigual, un split aleatorio simple no alcanza. ").bold = True
b.add_run("Como las canciones hit permanecen en el ranking una cantidad de semanas muy variable, fue "
          "necesario un algoritmo de asignación voraz (greedy bin-packing) para lograr que las proporciones "
          "de filas por subconjunto se acercaran al objetivo (60/15/25), en lugar de dejarlo librado al "
          "azar de qué canciones caen en cada grupo.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("Mejores métricas no siempre significan un mejor sistema, y peores métricas no siempre significan "
          "un sistema peor. ").bold = True
b.add_run("Lo importante es que la métrica sea confiable. En este caso, una evaluación metodológicamente "
          "más estricta arrojó métricas similares o incluso mejores que las de la Entrega 2, lo cual es el "
          "mejor escenario posible: rigurosidad sin sacrificar desempeño.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("Aislar la variable que se quiere comparar simplifica el análisis. ").bold = True
b.add_run("Generalizar el código para intercambiar la capa temporal (LSTM/GRU/Conv1D) sin modificar el "
          "resto de la arquitectura permitió atribuir las diferencias de desempeño específicamente al tipo "
          "de capa utilizada, en lugar de a cambios simultáneos en otras partes del modelo.")
b = doc.add_paragraph(style="List Bullet")
b.add_run("\"Usar el dataset completo\" no alcanza si el split sigue siendo por fila. ").bold = True
b.add_run("Se probó una variante que mantenía el dataset sin deduplicar (correcto) pero seleccionaba "
          "entrenamiento/validación/prueba con train_test_split estratificado por clase a nivel de fila, "
          "sin agrupar por canción. Al revisarla se encontró que 56 canciones quedaban repartidas entre "
          "entrenamiento y prueba, reproduciendo la misma fuga de información de la Entrega 2 sobre un "
          "dataset distinto, y con métricas aún más altas (F1 99,22%) precisamente por esa fuga. Esto "
          "confirmó, con un segundo caso independiente, que evitar la duplicación de canciones y evitar la "
          "fuga de información entre canciones son dos cosas distintas: conservar todas las filas es "
          "necesario pero no alcanza si el criterio de split no agrupa por la entidad que se repite.")

# ---------- Referencias ----------
doc.add_heading("Referencias", level=1)
refs = [
    "Dave, D. (2021). Billboard “The Hot 100” Songs. Kaggle. Disponible en: "
    "https://www.kaggle.com/datasets/dhruvildave/billboard-the-hot-100-songs",
    "Pandya, M. (2022). Spotify Tracks Dataset. Kaggle. Disponible en: "
    "https://www.kaggle.com/datasets/maharshipandya/-spotify-tracks-dataset",
    "Hochreiter, S., Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735–1780.",
    "Cho, K., van Merriënboer, B., Gulcehre, C., et al. (2014). Learning Phrase Representations using RNN "
    "Encoder-Decoder for Statistical Machine Translation. Proceedings of EMNLP 2014.",
    "LeCun, Y., Bengio, Y. (1995). Convolutional Networks for Images, Speech, and Time Series. "
    "The Handbook of Brain Theory and Neural Networks.",
    "TensorFlow Developers (2023). Keras: The high-level API for TensorFlow. TensorFlow Documentation. "
    "Disponible en: https://www.tensorflow.org/guide/keras",
    "TensorFlow Developers (2024). tf.keras.layers.GRU. TensorFlow API Documentation. Disponible en: "
    "https://www.tensorflow.org/api_docs/python/tf/keras/layers/GRU",
    "TensorFlow Developers (2024). tf.keras.layers.Conv1D. TensorFlow API Documentation. Disponible en: "
    "https://www.tensorflow.org/api_docs/python/tf/keras/layers/Conv1D",
    "Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning "
    "Research, 12, 2825-2830.",
    "pandas Development Team (s. f.). pandas Documentation. Disponible en: https://pandas.pydata.org/docs/",
]
for ref in refs:
    doc.add_paragraph(ref, style="List Number")

doc.add_paragraph()
doc.add_paragraph("Repositorio de código: https://github.com/mackhood/ia-avanzada")

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
